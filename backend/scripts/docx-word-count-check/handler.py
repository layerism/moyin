import json
import re
import sys
from pathlib import Path

from docx import Document


CJK_PATTERN = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
ALPHANUMERIC_PATTERN = re.compile(r"[A-Za-z0-9]+")


def count_words(text: str) -> int:
    return len(CJK_PATTERN.findall(text)) + len(ALPHANUMERIC_PATTERN.findall(text))


def document_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    seen_cells: set[int] = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                identity = id(cell._tc)
                if identity in seen_cells:
                    continue
                seen_cells.add(identity)
                parts.append(cell.text)
    return "\n".join(parts)


def run(payload: object) -> dict[str, object]:
    if not isinstance(payload, dict) or payload.get("schemaVersion") != "1.0":
        raise ValueError("不支持的输入协议版本")
    files = payload.get("files")
    if not isinstance(files, list) or len(files) != 1 or not isinstance(files[0], dict):
        raise ValueError("DOCX 字数审核只接受一个文件")
    material = files[0]
    file_id = material.get("id")
    path_value = material.get("path")
    extension = material.get("extension")
    if (
        not isinstance(file_id, str)
        or not file_id
        or not isinstance(path_value, str)
        or extension != ".docx"
    ):
        raise ValueError("DOCX 审核文件输入无效")
    context = payload.get("context")
    params = context.get("scriptParams") if isinstance(context, dict) else None
    minimum = params.get("minimumWordCount") if isinstance(params, dict) else None
    if (
        not isinstance(minimum, int)
        or isinstance(minimum, bool)
        or not 1 <= minimum <= 1_000_000
    ):
        raise ValueError("最低字数参数无效")
    path = Path(path_value)
    if path.suffix.lower() != ".docx" or not path.is_file():
        raise ValueError("DOCX 审核文件不存在")
    word_count = count_words(document_text(path))
    passed = word_count >= minimum
    message = f"当前 {word_count} 字，最低要求 {minimum} 字"
    issues = (
        []
        if passed
        else [
            {
                "fileId": file_id,
                "code": "WORD_COUNT_BELOW_MINIMUM",
                "message": message,
            }
        ]
    )
    return {
        "schemaVersion": "1.0",
        "passed": passed,
        "reason": f"文档字数符合要求：{message}" if passed else f"文档字数不足：{message}",
        "details": {
            "checkedFileCount": 1,
            "issues": issues,
            "wordCount": word_count,
            "minimumWordCount": minimum,
        },
    }


def main() -> None:
    payload = json.load(sys.stdin)
    json.dump(run(payload), sys.stdout, ensure_ascii=False)


if __name__ == "__main__":
    main()
